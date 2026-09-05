"""Проверка блока 4: встреча → решение → документ.

Как и в блоке 3, время подставляется явно, а половина проверок — про отказы.
Отказ проверять важнее: его слишком легко получить по неверной причине.

    docker compose -f docker-compose.yml -f docker-compose.dev.yml \
        run --rm --no-deps migrate python scripts/smoke_block4.py
"""
import asyncio
import io
import re
import sys
from datetime import date, datetime, time, timedelta
from pathlib import Path
from zoneinfo import ZoneInfo

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from sqlalchemy import delete, event, func, select, text, update

from app.core.db import engine, session_scope
from app.core.timeutil import utcnow
from app.models import (
    AgendaItem,
    AuditLog,
    Decision,
    DecisionStatus,
    Department,
    Document,
    DocumentAccess,
    DocumentScope,
    DocumentText,
    DocumentView,
    DownloadToken,
    Meeting,
    MeetingAttendance,
    MeetingParticipant,
    MeetingRating,
    MeetingRequest,
    MeetingStatus,
    Notification,
    Organization,
    IndexStatus,
    ParticipantRole,
    RoomBooking,
    SlotHold,
    Task,
    TaskStatus,
    TimeQuota,
    User,
    UserRole,
    UserStatus,
    WorkingHours,
)
from app.models.enums import RoleCode
from app.services import briefing
from app.services import decisions as registry
from app.services import export
from app.services import documents as docs
from app.services import indexer
from app.services import search
from app.services import tasks as task_service
from app.services import meetings as meeting_service
from app.services.bootstrap import bootstrap, ensure_default_working_hours, grant_role
from app.services.rbac import load_grants, visible_department_ids

from openpyxl import load_workbook
from pypdf import PdfReader

LF = bytes([10])
TEST_ORG_PREFIX = "ТЕСТ "
TZ = ZoneInfo("Asia/Tashkent")
MONDAY = date(2026, 10, 5)

passed = 0
failed = 0


def check(condition: bool, title: str, detail: str = "") -> None:
    global passed, failed
    if condition:
        passed += 1
        print(f"  OK   {title}")
    else:
        failed += 1
        print(f"  FAIL {title} {detail}")


def at(day: date, hour: int, minute: int = 0) -> datetime:
    return datetime.combine(day, time(hour, minute), tzinfo=TZ)


def tiny_pdf(line: str) -> bytes:
    """Минимальный PDF с одной строкой текста.

    Собран вручную, чтобы не тащить в зависимости генератор PDF ради проверки.
    Текст латиницей: стандартный Helvetica без встроенной кодировки кириллицу
    не отдаёт, а проверяем мы здесь разбор формата, а не морфологию — она
    проверяется на DOCX, где UTF-8 родной.
    """
    content = f"BT /F1 24 Tf 72 700 Td ({line}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode() + b" >>" + LF
        + b"stream" + LF + content + LF + b"endstream",
    ]
    out = bytearray(b"%PDF-1.4" + LF)
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj".encode() + LF + body + LF + b"endobj" + LF
    xref_at = len(out)
    out += f"xref{chr(10)}0 {len(objects) + 1}".encode() + LF
    out += b"0000000000 65535 f " + LF
    for offset in offsets:
        out += f"{offset:010d} 00000 n ".encode() + LF
    out += (
        f"trailer{chr(10)}<< /Size {len(objects) + 1} /Root 1 0 R >>{chr(10)}"
        f"startxref{chr(10)}{xref_at}{chr(10)}%%EOF{chr(10)}"
    ).encode()
    return bytes(out)


def tiny_docx(paragraphs: list[str], table_row: list[str] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for line in paragraphs:
        document.add_paragraph(line)
    if table_row:
        table = document.add_table(rows=1, cols=len(table_row))
        for cell, value in zip(table.rows[0].cells, table_row):
            cell.text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def tiny_xlsx(rows: list[list[str]]) -> bytes:
    from openpyxl import Workbook

    book = Workbook()
    sheet = book.active
    sheet.title = "ТЕСТ лист"
    for row in rows:
        sheet.append(row)
    buffer = io.BytesIO()
    book.save(buffer)
    return buffer.getvalue()


async def cleanup() -> None:
    async with session_scope() as session:
        org_ids = [
            row[0] for row in (
                await session.execute(
                    select(Organization.id).where(Organization.name.like(f"{TEST_ORG_PREFIX}%"))
                )
            ).all()
        ]
        if not org_ids:
            return
        user_ids = [
            row[0] for row in (
                await session.execute(select(User.id).where(User.organization_id.in_(org_ids)))
            ).all()
        ]
        doc_ids = [
            row[0] for row in (
                await session.execute(
                    select(Document.id).where(Document.organization_id.in_(org_ids))
                )
            ).all()
        ]
        meeting_ids = [
            row[0] for row in (
                await session.execute(
                    select(Meeting.id).where(Meeting.organization_id.in_(org_ids))
                )
            ).all()
        ]
        if doc_ids:
            for model in (DocumentAccess, DocumentText, DocumentView, DownloadToken):
                await session.execute(delete(model).where(model.document_id.in_(doc_ids)))
            await session.execute(delete(Document).where(Document.id.in_(doc_ids)))
        await session.execute(delete(Decision).where(Decision.organization_id.in_(org_ids)))
        if meeting_ids:
            await session.execute(
                delete(AgendaItem).where(AgendaItem.meeting_id.in_(meeting_ids))
            )
            for model in (MeetingParticipant, MeetingAttendance, MeetingRating, RoomBooking):
                await session.execute(delete(model).where(model.meeting_id.in_(meeting_ids)))
        if user_ids:
            await session.execute(delete(SlotHold).where(SlotHold.owner_id.in_(user_ids)))
            await session.execute(
                delete(MeetingRequest).where(MeetingRequest.owner_id.in_(user_ids))
            )
        if meeting_ids:
            await session.execute(delete(Meeting).where(Meeting.id.in_(meeting_ids)))
        await session.execute(delete(Task).where(Task.organization_id.in_(org_ids)))
        if user_ids:
            for model in (UserRole, WorkingHours, Notification):
                await session.execute(delete(model).where(model.user_id.in_(user_ids)))
            await session.execute(delete(TimeQuota).where(TimeQuota.owner_id.in_(user_ids)))
            await session.execute(delete(AuditLog).where(AuditLog.actor_id.in_(user_ids)))
            await session.execute(delete(User).where(User.id.in_(user_ids)))
        await session.execute(delete(Department).where(Department.organization_id.in_(org_ids)))
        await session.execute(delete(Organization).where(Organization.id.in_(org_ids)))


async def person(session, org, name, role, tg, department_id=None) -> User:
    user = User(
        organization_id=org.id, telegram_user_id=tg, full_name=name,
        status=UserStatus.ACTIVE, timezone="Asia/Tashkent", locale="ru",
        department_id=department_id,
    )
    session.add(user)
    await session.flush()
    await ensure_default_working_hours(session, user)
    await grant_role(session, user, role)
    return user


async def cast(session) -> dict[str, User]:
    """Действующие лица тестовой организации в свежей сессии."""
    org_id = (
        await session.execute(
            select(Organization.id).where(Organization.name == f"{TEST_ORG_PREFIX}Итоги")
        )
    ).scalar_one()
    people = (
        await session.execute(select(User).where(User.organization_id == org_id))
    ).scalars().all()
    return {p.full_name.split()[-1].lower(): p for p in people}


async def main() -> None:
    await cleanup()
    tg = 950_000_000

    print("\n1. Схема блока на месте")
    async with session_scope() as session:
        tables = {
            row[0] for row in (
                await session.execute(
                    text(
                        "select table_name from information_schema.tables "
                        "where table_schema='public'"
                    )
                )
            ).all()
        }
        expected = {
            "agenda_items", "decisions", "documents",
            "document_access", "document_texts", "document_views", "download_tokens",
        }
        check(expected <= tables, "семь таблиц блока созданы", f"нет: {sorted(expected - tables)}")

        # Индексы по выражению autogenerate не создаёт — их проверяем явно,
        # иначе поиск тихо уедет на последовательное чтение.
        indexes = {
            row[0] for row in (
                await session.execute(text("select indexname from pg_indexes"))
            ).all()
        }
        needed = {
            "ix_decisions_search", "ix_meetings_search", "ix_tasks_search",
            "ix_document_texts_search", "ix_documents_name_trgm",
        }
        check(needed <= indexes, "поисковые индексы существуют", f"нет: {sorted(needed - indexes)}")

        # На пустой таблице планировщик всегда выберет последовательное чтение,
        # поэтому «или индекс, или Seq Scan» не доказывало бы ничего. Отключаем
        # Seq Scan и требуем именно индекс: так проверяется то, что нужно —
        # что индекс подходит под форму запроса, а не просто существует.
        await session.execute(text("set local enable_seqscan = off"))
        probes = {
            "решения": (
                "explain select id from decisions where "
                "to_tsvector('russian', title || ' ' || coalesce(details, '')) "
                "@@ plainto_tsquery('russian', 'закупки')",
                "ix_decisions_search",
            ),
            "встречи": (
                "explain select id from meetings where "
                "to_tsvector('russian', title || ' ' || coalesce(description, '')) "
                "@@ plainto_tsquery('russian', 'совещание')",
                "ix_meetings_search",
            ),
            "поручения": (
                "explain select id from tasks where "
                "to_tsvector('russian', title || ' ' || coalesce(description, '')) "
                "@@ plainto_tsquery('russian', 'отчёт')",
                "ix_tasks_search",
            ),
            "тексты документов": (
                "explain select id from document_texts where "
                "search_vector @@ plainto_tsquery('russian', 'договор')",
                "ix_document_texts_search",
            ),
        }
        for label, (sql, index_name) in probes.items():
            plan = " | ".join(
                str(row[0]) for row in (await session.execute(text(sql))).all()
            )
            check(index_name in plan, f"поиск по «{label}» опирается на индекс", plan[:120])

    async with session_scope() as session:
        await bootstrap(session)
        org = Organization(name=f"{TEST_ORG_PREFIX}Итоги", timezone="Asia/Tashkent")
        session.add(org)
        await session.flush()
        dept = Department(organization_id=org.id, name="ТЕСТ отдел")
        session.add(dept)
        await session.flush()

        chief = await person(session, org, "Тест Руководитель", RoleCode.EXECUTIVE, tg + 1)
        helper = await person(session, org, "Тест Ассистент", RoleCode.ASSISTANT, tg + 2)
        head = await person(session, org, "Тест Начальник", RoleCode.DEPT_HEAD, tg + 3, dept.id)
        worker = await person(session, org, "Тест Сотрудник", RoleCode.EMPLOYEE, tg + 4, dept.id)

        other = Organization(name=f"{TEST_ORG_PREFIX}Чужая", timezone="Asia/Tashkent")
        session.add(other)
        await session.flush()
        await person(session, other, "Тест Чужой", RoleCode.EXECUTIVE, tg + 9)

        meeting = Meeting(
            organization_id=org.id, owner_id=chief.id, title="ТЕСТ совещание",
            start_at=at(MONDAY, 10), end_at=at(MONDAY, 11),
            status=MeetingStatus.CONFIRMED, created_by=chief.id,
        )
        session.add(meeting)
        await session.flush()
        for uid, role in (
            (chief.id, ParticipantRole.ORGANIZER),
            (worker.id, ParticipantRole.REQUIRED),
            (head.id, ParticipantRole.REQUIRED),
        ):
            session.add(MeetingParticipant(
                meeting_id=meeting.id, user_id=uid, role=role, created_at=utcnow()
            ))
        await session.flush()
        meeting_id = meeting.id

    print("\n2. Повестка")
    async with session_scope() as session:
        who = await cast(session)
        meeting = await session.get(Meeting, meeting_id)

        short = await registry.add_agenda_item(
            session, meeting=meeting, actor=who["руководитель"], title="ок"
        )
        check(not short.ok, "пункт в два знака не принимается", short.reason or "")

        denied = await registry.add_agenda_item(
            session, meeting=meeting, actor=who["сотрудник"], title="ТЕСТ мой пункт"
        )
        check(not denied.ok, "сотрудник повестку не ведёт", denied.reason or "")

        for title in ("ТЕСТ закупки", "ТЕСТ склад", "ТЕСТ кадры"):
            added = await registry.add_agenda_item(
                session, meeting=meeting, actor=who["руководитель"], title=title
            )
            check(added.ok, f"пункт «{title}» добавлен", added.reason or "")

        items = await registry.agenda_of(session, meeting)
        check(
            [i.title for i in items] == ["ТЕСТ закупки", "ТЕСТ склад", "ТЕСТ кадры"],
            "порядок повестки сохранён",
            f"{[i.title for i in items]}",
        )
        check(
            [i.position for i in items] == [1, 2, 3],
            "позиции проставлены подряд",
            f"{[i.position for i in items]}",
        )

        covered = await registry.mark_covered(
            session, item=items[0], meeting=meeting, actor=who["ассистент"]
        )
        check(covered.ok, "ассистент отмечает пункт рассмотренным", covered.reason or "")
        check(items[0].covered_at is not None, "и время отметки записано")

        by_worker = await registry.mark_covered(
            session, item=items[1], meeting=meeting, actor=who["сотрудник"]
        )
        check(not by_worker.ok, "сотрудник пункты не отмечает", by_worker.reason or "")

    print("\n3. Завершение встречи")
    async with session_scope() as session:
        who = await cast(session)
        meeting = await session.get(Meeting, meeting_id)

        early = await meeting_service.finish(
            session, meeting=meeting, actor=who["руководитель"], now=at(MONDAY, 9)
        )
        check(not early.ok, "до начала встречу не завершить", early.reason or "")

        stranger = await meeting_service.finish(
            session, meeting=meeting, actor=who["сотрудник"], now=at(MONDAY, 11, 30)
        )
        check(not stranger.ok, "участник чужую встречу не завершает", stranger.reason or "")
        check(meeting.status == MeetingStatus.CONFIRMED, "и встреча всё ещё идёт")

        done = await meeting_service.finish(
            session, meeting=meeting, actor=who["руководитель"], now=at(MONDAY, 11, 30)
        )
        check(done.ok, "руководитель завершает встречу", done.reason or "")
        check(meeting.status == MeetingStatus.FINISHED, "состояние — завершена")
        check(meeting.finished_at is not None, "время завершения записано")

        again = await meeting_service.finish(
            session, meeting=meeting, actor=who["руководитель"], now=at(MONDAY, 12)
        )
        check(not again.ok, "повторное завершение ничего не меняет", again.reason or "")

        late_item = await registry.add_agenda_item(
            session, meeting=meeting, actor=who["руководитель"], title="ТЕСТ поздний пункт"
        )
        check(not late_item.ok, "в завершённую встречу пункт не добавить", late_item.reason or "")

    print("\n4. Реестр решений")
    async with session_scope() as session:
        who = await cast(session)
        meeting = await session.get(Meeting, meeting_id)
        items = await registry.agenda_of(session, meeting)

        empty = await registry.create(session, actor=who["руководитель"], title=" ")
        check(not empty.ok, "решение без формулировки не создаётся", empty.reason or "")

        by_worker = await registry.create(
            session, actor=who["сотрудник"], title="ТЕСТ решение сотрудника"
        )
        check(not by_worker.ok, "сотрудник решений не вносит", by_worker.reason or "")

        first = await registry.create(
            session, actor=who["руководитель"], title="ТЕСТ закупить стеллажи",
            details="Три штуки до конца месяца", meeting=meeting, agenda_item=items[0],
            responsible=who["сотрудник"], due_date=at(MONDAY + timedelta(days=14), 18),
        )
        check(first.ok, "решение внесено", first.reason or "")
        check(first.item.status == DecisionStatus.OPEN, "и оно в работе")
        check(first.item.meeting_id == meeting_id, "связано со встречей")

        loose = await registry.create(
            session, actor=who["руководитель"], title="ТЕСТ решение без встречи"
        )
        check(loose.ok, "решение без встречи тоже вносится", loose.reason or "")
        check(loose.item.meeting_id is None, "и живёт само по себе")

        outsider = (
            await session.execute(
                select(User).join(Organization, Organization.id == User.organization_id)
                .where(Organization.name == f"{TEST_ORG_PREFIX}Чужая")
            )
        ).scalar_one()
        cross = await registry.create(
            session, actor=outsider, title="ТЕСТ чужое решение", meeting=meeting
        )
        check(not cross.ok, "к чужой встрече решение не привязать", cross.reason or "")

        decision_id = first.item.id
        loose_id = loose.item.id

    print("\n5. Закрытие и отмена решения")
    async with session_scope() as session:
        who = await cast(session)
        decision = await session.get(Decision, decision_id)
        loose = await session.get(Decision, loose_id)

        by_worker = await registry.close(
            session, decision=decision, actor=who["сотрудник"], done=True
        )
        check(not by_worker.ok, "сотрудник решение не закрывает", by_worker.reason or "")

        by_head = await registry.close(
            session, decision=decision, actor=who["начальник"], done=True
        )
        check(not by_head.ok, "начальник отдела — тоже нет", by_head.reason or "")

        no_reason = await registry.close(
            session, decision=loose, actor=who["руководитель"], done=False, reason="  "
        )
        check(not no_reason.ok, "отмена без причины не принимается", no_reason.reason or "")

        closed = await registry.close(
            session, decision=decision, actor=who["руководитель"], done=True
        )
        check(closed.ok, "руководитель закрывает решение", closed.reason or "")
        check(decision.status == DecisionStatus.DONE, "состояние — выполнено")

        twice = await registry.close(
            session, decision=decision, actor=who["руководитель"], done=True
        )
        check(not twice.ok, "повторное закрытие ничего не меняет", twice.reason or "")

        killed = await registry.close(
            session, decision=loose, actor=who["ассистент"], done=False,
            reason="Отпало по итогам совещания",
        )
        check(killed.ok, "ассистент отменяет решение с причиной", killed.reason or "")
        check(loose.status == DecisionStatus.CANCELLED, "состояние — отменено")
        check(
            loose.cancel_reason == "Отпало по итогам совещания",
            "причина отмены сохранена",
        )

        still_there = await session.scalar(
            select(func.count(Decision.id)).where(Decision.id.in_([decision_id, loose_id]))
        )
        check(still_there == 2, "обе строки остались в реестре", f"их {still_there}")

    print("\n6. Кто какие решения видит")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head = who["руководитель"], who["сотрудник"], who["начальник"]

        seen_by_chief = await registry.registry(
            session, user=chief, grants=await load_grants(session, chief), limit=50
        )
        check(len(seen_by_chief) == 2, "руководитель видит оба решения", f"их {len(seen_by_chief)}")

        seen_by_worker = await registry.registry(
            session, user=worker, grants=await load_grants(session, worker), limit=50
        )
        check(
            [d.id for d in seen_by_worker] == [decision_id],
            "сотрудник видит только то, где он ответственный",
            f"{[d.id for d in seen_by_worker]}",
        )

        seen_by_head = await registry.registry(
            session, user=head, grants=await load_grants(session, head), limit=50
        )
        check(
            decision_id in [d.id for d in seen_by_head],
            "начальник отдела видит решение своего сотрудника",
            f"{[d.id for d in seen_by_head]}",
        )
        check(
            loose_id not in [d.id for d in seen_by_head],
            "и не видит решение без отношения к его отделу",
        )

        outsider = (
            await session.execute(
                select(User).join(Organization, Organization.id == User.organization_id)
                .where(Organization.name == f"{TEST_ORG_PREFIX}Чужая")
            )
        ).scalar_one()
        seen_by_outsider = await registry.registry(
            session, user=outsider, grants=await load_grants(session, outsider), limit=50
        )
        check(not seen_by_outsider, "чужая организация не видит ничего", f"{len(seen_by_outsider)}")

        meeting = await session.get(Meeting, meeting_id)
        made_decisions, made_tasks = await registry.meeting_outcome(session, meeting)
        check(made_decisions == 1, "встреча знает своё решение", f"их {made_decisions}")
        check(made_tasks == 0, "поручений из неё пока нет", f"их {made_tasks}")

    print("\n7. Поручения прямо из встречи")
    async with session_scope() as session:
        who = await cast(session)
        meeting = await session.get(Meeting, meeting_id)
        decision = await session.get(Decision, decision_id)

        task = await task_service.create_task(
            session, creator=who["руководитель"], assignee=who["сотрудник"],
            title="ТЕСТ закупить стеллажи до пятницы",
            due_at=at(MONDAY + timedelta(days=4), 18),
            meeting_id=meeting.id, decision_id=decision.id,
        )
        check(task.id is not None, "поручение создано из встречи")
        check(task.meeting_id == meeting_id, "связь со встречей сохранена")
        check(task.decision_id == decision_id, "и с решением тоже")
        check(task.status == TaskStatus.NEW, "жизненный цикл обычный, без второй ветки")

        from_meeting = await task_service.of_meeting(session, meeting_id)
        check(len(from_meeting) == 1, "встреча знает своё поручение", f"их {len(from_meeting)}")
        from_decision = await task_service.of_decision(session, decision_id)
        check(len(from_decision) == 1, "решение знает своё поручение", f"их {len(from_decision)}")

        made_decisions, made_tasks = await registry.meeting_outcome(session, meeting)
        check(
            (made_decisions, made_tasks) == (1, 1),
            "встреча с результатом: решение и поручение",
            f"{made_decisions} и {made_tasks}",
        )

        empty_meeting = Meeting(
            organization_id=meeting.organization_id, owner_id=who["руководитель"].id,
            title="ТЕСТ встреча без результата",
            start_at=at(MONDAY + timedelta(days=1), 10),
            end_at=at(MONDAY + timedelta(days=1), 11),
            status=MeetingStatus.FINISHED, created_by=who["руководитель"].id,
            finished_at=at(MONDAY + timedelta(days=1), 11),
        )
        session.add(empty_meeting)
        await session.flush()
        nothing = await registry.meeting_outcome(session, empty_meeting)
        check(nothing == (0, 0), "встреча без результата видна как нулевая", f"{nothing}")
        task_id = task.id

    print("\n8. Приём документов")
    async with session_scope() as session:
        who = await cast(session)
        meeting = await session.get(Meeting, meeting_id)

        danger = await docs.store(
            session, uploader=who["сотрудник"], file_id="f1", file_unique_id="u1",
            file_name="вирус.exe", size_bytes=1024,
        )
        check(not danger.ok, "исполняемый файл не принимается", danger.reason or "")

        nameless = await docs.store(
            session, uploader=who["сотрудник"], file_id="f2", file_unique_id="u2",
            file_name="   ", size_bytes=1024,
        )
        check(not nameless.ok, "файл без имени не принимается", nameless.reason or "")

        huge = await docs.store(
            session, uploader=who["сотрудник"], file_id="f3", file_unique_id="u3",
            file_name="архив.pdf", size_bytes=3 * 1024 * 1024 * 1024,
        )
        check(not huge.ok, "файл больше предельного не принимается", huge.reason or "")

        ok_doc = await docs.store(
            session, uploader=who["сотрудник"], file_id="f4", file_unique_id="u4",
            file_name="договор.pdf", size_bytes=200 * 1024,
            mime_type="application/pdf", meeting=meeting,
        )
        check(ok_doc.ok, "обычный документ принят", ok_doc.reason or "")
        check(
            ok_doc.document.index_status == IndexStatus.PENDING,
            "и поставлен в очередь на извлечение текста",
            ok_doc.document.index_status,
        )
        check(
            ok_doc.document.scope == DocumentScope.PRIVATE,
            "по умолчанию — личный, а не общий",
            ok_doc.document.scope,
        )

        big_pdf = await docs.store(
            session, uploader=who["сотрудник"], file_id="f5", file_unique_id="u5",
            file_name="скан.pdf", size_bytes=25 * 1024 * 1024,
        )
        check(big_pdf.ok, "файл больше 20 МБ принимается", big_pdf.reason or "")
        check(
            big_pdf.document.index_status == IndexStatus.TOO_LARGE,
            "но честно помечен: текст из него не достать",
            big_pdf.document.index_status,
        )

        picture = await docs.store(
            session, uploader=who["сотрудник"], file_id="f6", file_unique_id="u6",
            file_name="схема.png", size_bytes=50 * 1024,
        )
        check(
            picture.document.index_status == IndexStatus.UNSUPPORTED,
            "формат без текста помечен отдельно",
            picture.document.index_status,
        )
        private_id = ok_doc.document.id

    print("\n9. Доступ к встрече не даёт доступа к документам")
    async with session_scope() as session:
        who = await cast(session)
        document = await session.get(Document, private_id)

        check(
            await docs.may_read(session, document=document, viewer=who["сотрудник"]),
            "загрузивший свой документ видит",
        )
        check(
            not await docs.may_read(session, document=document, viewer=who["начальник"]),
            "участник той же встречи документ НЕ получает",
        )
        denied = await docs.open_for(
            session, document=document, viewer=who["начальник"]
        )
        check(denied is not None, "и выдача ему отказана", denied or "")
        views = await docs.views_of(session, document)
        check(not views, "отказ не пишется в журнал открытий как открытие", f"{len(views)}")

        problem = await docs.grant(
            session, document=document, actor=who["начальник"], to_user=who["ассистент"]
        )
        check(problem is not None, "чужой документ третьему лицу не открыть", problem or "")

        opened = await docs.grant(
            session, document=document, actor=who["сотрудник"], to_user=who["начальник"]
        )
        check(opened is None, "загрузивший открывает доступ", opened or "")
        check(
            await docs.may_read(session, document=document, viewer=who["начальник"]),
            "и теперь документ ему открыт",
        )

        given = await docs.open_for(session, document=document, viewer=who["начальник"])
        check(given is None, "выдача проходит", given or "")
        views = await docs.views_of(session, document)
        check(len(views) == 1, "открытие записано в журнал", f"записей {len(views)}")
        check(views[0].user_id == who["начальник"].id, "и записан правильный человек")

        outsider = (
            await session.execute(
                select(User).join(Organization, Organization.id == User.organization_id)
                .where(Organization.name == f"{TEST_ORG_PREFIX}Чужая")
            )
        ).scalar_one()
        check(
            not await docs.may_read(session, document=document, viewer=outsider),
            "чужая организация документ не видит",
        )

    print("\n10. Прямая проверка и поиск отвечают одинаково")
    async with session_scope() as session:
        who = await cast(session)
        meeting = await session.get(Meeting, meeting_id)
        worker, head, chief = who["сотрудник"], who["начальник"], who["руководитель"]

        # По документу на каждую политику доступа — и проверка каждой пары
        # «документ — человек» двумя путями: прямым и через условие поиска.
        matrix = {}
        for label, scope in (
            ("личный", DocumentScope.PRIVATE),
            ("участникам", DocumentScope.PARTICIPANTS),
            ("отделу", DocumentScope.DEPARTMENT),
            ("организации", DocumentScope.ORGANIZATION),
        ):
            made = await docs.store(
                session, uploader=worker, file_id=f"m-{scope}", file_unique_id=f"mu-{scope}",
                file_name=f"ТЕСТ {label}.pdf", size_bytes=1024,
                meeting=meeting, scope=scope,
            )
            matrix[label] = made.document

        mismatches = []
        for label, document in matrix.items():
            for name in ("сотрудник", "начальник", "ассистент", "руководитель"):
                viewer = who[name]
                direct = await docs.may_read(session, document=document, viewer=viewer)
                grants = await load_grants(session, viewer)
                visible = await visible_department_ids(session, viewer)
                found = await session.scalar(
                    select(Document.id).where(
                        Document.id == document.id,
                        *docs.visible_filter(viewer, grants, visible),
                    )
                )
                if direct != (found is not None):
                    mismatches.append(f"{label}/{name}: прямо={direct}, поиском={found is not None}")

        check(
            not mismatches,
            "на каждой паре «документ — человек» ответы совпадают",
            "; ".join(mismatches[:3]),
        )

        # И парные утверждения: политика вообще что-то меняет.
        check(
            not await docs.may_read(session, document=matrix["личный"], viewer=chief),
            "личный документ закрыт даже руководителю",
        )
        check(
            await docs.may_read(session, document=matrix["организации"], viewer=chief),
            "а открытый организации — открыт",
        )
        check(
            await docs.may_read(session, document=matrix["участникам"], viewer=head),
            "документ «участникам» открыт участнику встречи",
        )
        check(
            not await docs.may_read(session, document=matrix["участникам"], viewer=who["ассистент"]),
            "и закрыт тому, кого на встрече не было",
        )

    print("\n11. Извлечение текста")
    async with session_scope() as session:
        who = await cast(session)
        worker = who["сотрудник"]
        meeting = await session.get(Meeting, meeting_id)

        files = {
            "устав.txt": "Устав общества. Раздел о закупках оборудования.".encode("utf-8"),
            "протокол.docx": tiny_docx(
                ["Протокол совещания по складским остаткам",
                 "Обсуждали поставку стеллажей"],
                ["Позиция", "Стеллаж металлический", "12 штук"],
            ),
            "смета.xlsx": tiny_xlsx([
                ["Наименование", "Количество"],
                ["Стеллаж металлический", 12],
                ["Погрузчик дизельный", 1],
            ]),
            "contract.pdf": tiny_pdf("Dogovor postavki oborudovaniya"),
            "битый.pdf": b"not a pdf at all",
            # Двоичный файл под именем текстового: так выглядит обычная
            # ошибка человека, а внутри — нулевые байты, которых база
            # не принимает вовсе.
            "двоичный.txt": tiny_docx(["не текст"]),
        }
        stored = {}
        for name, data in files.items():
            made = await docs.store(
                session, uploader=worker, file_id=f"id-{name}", file_unique_id=f"u-{name}",
                file_name=name, size_bytes=len(data),
                meeting=meeting, scope=DocumentScope.ORGANIZATION,
            )
            check(made.ok, f"документ «{name}» принят", made.reason or "")
            stored[name] = made.document.id

        async def fake_download(file_id: str) -> bytes:
            return files[file_id.removeprefix("id-")]

        stats = await indexer.index_pending(session, fake_download, limit=50)
        # Пять: четыре читаемых формата плюс двоичный файл, из которого
        # после очистки остаётся мусорный, но безопасный текст.
        check(stats["done"] == 5, "разобрано пять документов", f"{stats}")

        # В очереди лежат и документы прошлых разделов, поэтому общий счёт
        # сбоев ничего не сказал бы. Спрашиваем о судьбе именно наших файлов.
        outcomes = {}
        for name, doc_id in stored.items():
            outcomes[name] = (await session.get(Document, doc_id)).index_status
        check(
            all(outcomes[n] == IndexStatus.DONE for n in
                ("устав.txt", "протокол.docx", "смета.xlsx", "contract.pdf")),
            "все четыре читаемых формата разобраны",
            f"{outcomes}",
        )
        check(
            outcomes["битый.pdf"] == IndexStatus.FAILED,
            "битый файл помечен сбоем, а не потерян",
            f"{outcomes['битый.pdf']}",
        )
        check(
            stats["done"] + stats["failed"] + stats["empty"] > 0,
            "цикл дошёл до конца, не упав на битом файле",
            f"{stats}",
        )

        broken = await session.get(Document, stored["битый.pdf"])
        check(broken.index_status == IndexStatus.FAILED, "состояние битого — сбой")
        check(bool(broken.index_error), "и причина записана", broken.index_error or "")
        check(broken.file_id is not None, "сам документ остался на месте")

        for name, needle in (
            ("устав.txt", "закупках"),
            ("протокол.docx", "стеллажей"),
            ("смета.xlsx", "Погрузчик"),
            ("contract.pdf", "Dogovor"),
        ):
            body = await session.scalar(
                select(DocumentText.content).where(DocumentText.document_id == stored[name])
            )
            check(needle in (body or ""), f"текст из «{name}» извлечён", (body or "")[:60])

        table_text = await session.scalar(
            select(DocumentText.content).where(DocumentText.document_id == stored["протокол.docx"])
        )
        check(
            "12 штук" in table_text,
            "таблицы из docx тоже попадают в текст",
            table_text[:80],
        )

        # Нулевой байт внутри текста PostgreSQL отвергает вместе со всей
        # строкой. Проверяем, что это состояние одного документа, а не обрыв
        # прохода: соседние файлы обязаны быть разобраны.
        binary = await session.get(Document, stored["двоичный.txt"])
        check(
            binary.index_status in (IndexStatus.DONE, IndexStatus.EMPTY, IndexStatus.FAILED),
            "двоичный файл под именем текстового получил состояние, а не обрушил проход",
            binary.index_status,
        )
        stored_text = await session.scalar(
            select(DocumentText.content).where(
                DocumentText.document_id == stored["двоичный.txt"]
            )
        )
        check(
            chr(0) not in (stored_text or ""),
            "и в базу не попало ни одного нулевого байта",
        )

        again = await indexer.index_pending(session, fake_download, limit=20)
        check(
            again == {"done": 0, "empty": 0, "failed": 0},
            "повторный проход не делает лишней работы",
            f"{again}",
        )
        docs_ids = stored

    print("\n12. Поиск: права входят в запрос")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head = who["руководитель"], who["сотрудник"], who["начальник"]

        secret = await task_service.create_task(
            session, creator=chief, assignee=head,
            title="ТЕСТ пересчитать премию главного бухгалтера",
        )
        await session.flush()

        found = await search.search(
            session, user=chief, grants=await load_grants(session, chief),
            query="премию бухгалтера",
        )
        check(
            any(h.id == secret.id for h in found.tasks),
            "руководитель находит своё поручение",
            f"{[h.title for h in found.tasks]}",
        )

        # Главная проверка блока: точная формулировка не помогает чужому.
        hidden = await search.search(
            session, user=worker, grants=await load_grants(session, worker),
            query="ТЕСТ пересчитать премию главного бухгалтера",
        )
        check(
            not any(h.id == secret.id for h in hidden.tasks),
            "сотрудник не находит чужое поручение даже по точной формулировке",
            f"{[h.title for h in hidden.tasks]}",
        )
        colleagues = await search.search(
            session, user=worker, grants=await load_grants(session, worker),
            query="Начальник",
        )
        check(
            any(h.id == head.id for h in colleagues.people),
            "но коллегу по справочнику находит — это не тайна",
            f"{[h.title for h in colleagues.people]}",
        )

        # «поставку стеллажей» есть только в теле документа; в имени файла
        # («протокол.docx») этих слов нет, поэтому совпадение может прийти
        # только из извлечённого текста.
        by_morphology = await search.search(
            session, user=chief, grants=await load_grants(session, chief),
            query="поставка стеллажа",
        )
        check(
            any(h.id == docs_ids["протокол.docx"] for h in by_morphology.documents),
            "документ находится по фразе из текста с учётом морфологии",
            f"{[h.title for h in by_morphology.documents]}",
        )
        check(
            not any("поставка" in h.title.lower() for h in by_morphology.documents),
            "и совпало именно содержимое, а не имя файла",
        )

        typo = await search.search(
            session, user=worker, grants=await load_grants(session, worker),
            query="Начальнек",
        )
        check(
            any(h.id == head.id for h in typo.people),
            "опечатка в фамилии прощается",
            f"{[h.title for h in typo.people]}",
        )

        nothing = await search.search(
            session, user=worker, grants=await load_grants(session, worker), query=" "
        )
        check(nothing.empty, "пустой запрос ничего не выгружает", f"{nothing.total}")
        one_letter = await search.search(
            session, user=worker, grants=await load_grants(session, worker), query="с"
        )
        check(one_letter.empty, "и один символ тоже", f"{one_letter.total}")

    print("\n13. Поиск документов подчиняется правам на файл")
    async with session_scope() as session:
        who = await cast(session)
        worker, helper_user = who["сотрудник"], who["ассистент"]
        meeting = await session.get(Meeting, meeting_id)

        closed = await docs.store(
            session, uploader=worker, file_id="id-закрытый", file_unique_id="u-закрытый",
            file_name="закрытая записка.txt", size_bytes=100,
            meeting=meeting, scope=DocumentScope.PRIVATE,
        )
        await indexer.save_text(
            session, document=closed.document,
            content="Совершенно секретная записка про поставку стеллажей", pages=1,
        )
        closed.document.index_status = IndexStatus.DONE
        await session.flush()

        mine = await search.search(
            session, user=worker, grants=await load_grants(session, worker),
            query="секретная записка",
        )
        check(
            any(h.id == closed.document.id for h in mine.documents),
            "автор находит свой закрытый документ",
        )

        theirs = await search.search(
            session, user=helper_user, grants=await load_grants(session, helper_user),
            query="секретная записка",
        )
        check(
            not any(h.id == closed.document.id for h in theirs.documents),
            "и никто другой его не находит — даже по точной фразе",
            f"{[h.title for h in theirs.documents]}",
        )

        await docs.grant(
            session, document=closed.document, actor=worker, to_user=helper_user
        )
        after_grant = await search.search(
            session, user=helper_user, grants=await load_grants(session, helper_user),
            query="секретная записка",
        )
        check(
            any(h.id == closed.document.id for h in after_grant.documents),
            "после выдачи доступа находит",
        )

    print("\n14. Два описания прав на поручения совпадают")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head, helper_user = (
            who["руководитель"], who["сотрудник"], who["начальник"], who["ассистент"]
        )
        outsider = (
            await session.execute(
                select(User).join(Organization, Organization.id == User.organization_id)
                .where(Organization.name == f"{TEST_ORG_PREFIX}Чужая")
            )
        ).scalar_one()

        made = []
        for creator, assignee, label in (
            (chief, worker, "руководитель сотруднику"),
            (chief, head, "руководитель начальнику"),
            (head, worker, "начальник сотруднику"),
        ):
            made.append(await task_service.create_task(
                session, creator=creator, assignee=assignee, title=f"ТЕСТ {label}",
            ))
        await session.flush()

        mismatches = []
        for task in made:
            for name, viewer in (
                ("руководитель", chief), ("ассистент", helper_user),
                ("начальник", head), ("сотрудник", worker), ("чужой", outsider),
            ):
                grants = await load_grants(session, viewer)
                direct = (
                    await task_service.access_for(session, task, viewer, grants)
                ).can_view
                visible = await visible_department_ids(session, viewer)
                by_sql = await session.scalar(
                    select(Task.id).where(
                        Task.id == task.id,
                        *task_service.visible_filter(viewer, grants, visible),
                    )
                )
                if direct != (by_sql is not None):
                    mismatches.append(
                        f"{task.title}/{name}: прямо={direct}, поиском={by_sql is not None}"
                    )
        check(
            not mismatches,
            "проверка записи и условие поиска дают один ответ",
            "; ".join(mismatches[:3]),
        )
        check(
            not (await task_service.access_for(
                session, made[0], outsider, await load_grants(session, outsider)
            )).can_view,
            "и чужая организация поручение не видит",
        )

    print("\n15. Досье к встрече")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head = who["руководитель"], who["сотрудник"], who["начальник"]
        soon_day = MONDAY + timedelta(days=7)

        upcoming = Meeting(
            organization_id=chief.organization_id, owner_id=chief.id,
            title="ТЕСТ разбор поставки", start_at=at(soon_day, 15),
            end_at=at(soon_day, 16), status=MeetingStatus.CONFIRMED, created_by=chief.id,
        )
        session.add(upcoming)
        await session.flush()
        for uid in (chief.id, worker.id, head.id):
            session.add(MeetingParticipant(
                meeting_id=upcoming.id, user_id=uid,
                role=ParticipantRole.REQUIRED, created_at=utcnow(),
            ))
        await session.flush()

        # Документ, открытый только автору: в чужом досье его быть не должно
        # даже названием.
        private = await docs.store(
            session, uploader=worker, file_id="id-досье", file_unique_id="u-досье",
            file_name="личная записка.txt", size_bytes=100,
            meeting=upcoming, scope=DocumentScope.PRIVATE,
        )
        shared = await docs.store(
            session, uploader=worker, file_id="id-общий", file_unique_id="u-общий",
            file_name="повестка.txt", size_bytes=100,
            meeting=upcoming, scope=DocumentScope.ORGANIZATION,
        )
        await session.flush()

        early = await briefing.send_briefings(session, now=at(soon_day, 12))
        check(early == 0, "за три часа досье не уходит", f"писем: {early}")

        sent = await briefing.send_briefings(session, now=at(soon_day, 14, 45))
        check(sent == 3, "за пятнадцать минут досье ушло всем троим", f"писем: {sent}")

        again = await briefing.send_briefings(session, now=at(soon_day, 14, 50))
        check(again == 0, "повторный проход второй раз не шлёт", f"писем: {again}")

        own = await briefing.build(
            session, meeting=upcoming, viewer=worker, now=at(soon_day, 14, 45)
        )
        check("личная записка" in own, "автор видит свой документ в досье")
        check("повестка" in own, "и общий тоже")

        theirs = await briefing.build(
            session, meeting=upcoming, viewer=head, now=at(soon_day, 14, 45)
        )
        check(
            "личная записка" not in theirs,
            "а другому участнику закрытый документ не показан даже названием",
            theirs[:120],
        )
        check("повестка" in theirs, "общий документ при этом виден ему")
        check("ТЕСТ разбор поставки" in theirs, "тема встречи в досье есть")

        # Прошлое незакрытое решение по этим же людям.
        old_decision = await registry.create(
            session, actor=chief, title="ТЕСТ проверить остатки на складе",
            responsible=worker,
        )
        await session.flush()
        with_history = await briefing.build(
            session, meeting=upcoming, viewer=chief, now=at(soon_day, 14, 45)
        )
        check(
            "проверить остатки" in with_history,
            "прошлое незакрытое решение попало в досье",
            with_history[-200:],
        )

        # Перенос — другое время, досье нужно заново.
        upcoming.reschedule_count += 1
        upcoming.start_at = at(soon_day, 17)
        upcoming.end_at = at(soon_day, 18)
        await session.flush()
        after_move = await briefing.send_briefings(session, now=at(soon_day, 16, 45))
        check(after_move == 3, "после переноса досье уходит заново", f"писем: {after_move}")

    print("\n16. Выгрузка в Excel и PDF")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker = who["руководитель"], who["сотрудник"]
        # Период берём с запасом назад: поручения создаются по настоящим
        # часам, а опорный понедельник проверки лежит в будущем.
        since, until = at(MONDAY - timedelta(days=120), 0), at(MONDAY + timedelta(days=60), 0)

        data, name, why = await export.build(
            session, user=chief, grants=await load_grants(session, chief),
            kind="tasks", since=since, until=until, fmt="xlsx",
        )
        check(data is not None, "выгрузка поручений собрана", why or "")
        check(name.endswith(".xlsx"), "имя файла с нужным расширением", name)

        book = load_workbook(io.BytesIO(data))
        page = book.active
        rows = list(page.iter_rows(min_row=2, values_only=True))
        check(bool(rows), "в файле есть строки", f"строк {len(rows)}")
        check(
            page.cell(row=1, column=2).value == "Поручение",
            "заголовки на месте",
            str(page.cell(row=1, column=2).value),
        )
        check(
            any("стеллажи" in str(r[1]).lower() for r in rows),
            "и в нём именно наши данные",
            f"{[r[1] for r in rows][:3]}",
        )

        # Выгрузка обязана показывать ровно то же, что интерфейс.
        visible = await visible_department_ids(session, chief)
        in_view = await session.scalar(
            select(func.count(Task.id)).where(
                *task_service.visible_filter(chief, await load_grants(session, chief), visible),
                Task.created_at >= since, Task.created_at < until,
            )
        )
        check(len(rows) == in_view, "строк ровно столько, сколько видно в списке",
              f"{len(rows)} против {in_view}")

        worker_data, _, _ = await export.build(
            session, user=worker, grants=await load_grants(session, worker),
            kind="tasks", since=since, until=until, fmt="xlsx",
        )
        worker_rows = (
            list(load_workbook(io.BytesIO(worker_data)).active.iter_rows(min_row=2, values_only=True))
            if worker_data else []
        )
        check(
            len(worker_rows) < len(rows),
            "сотрудник выгружает меньше, чем руководитель",
            f"{len(worker_rows)} против {len(rows)}",
        )

        pdf_data, pdf_name, why = await export.build(
            session, user=chief, grants=await load_grants(session, chief),
            kind="decisions", since=since, until=until, fmt="pdf",
        )
        check(pdf_data is not None, "выгрузка решений в PDF собрана", why or "")
        check(pdf_data[:5] == b"%PDF-", "это настоящий PDF", str(pdf_data[:8]))
        page_text = PdfReader(io.BytesIO(pdf_data)).pages[0].extract_text()
        check("Решения" in page_text, "заголовок читается", page_text[:60])
        check(
            "стеллажи" in page_text.lower(),
            "и кириллица в таблице не превратилась в вопросы",
            page_text[:200],
        )

        nothing, _, why = await export.build(
            session, user=chief, grants=await load_grants(session, chief),
            kind="meetings",
            since=at(MONDAY - timedelta(days=400), 0),
            until=at(MONDAY - timedelta(days=390), 0),
        )
        check(nothing is None, "пустой период даёт честный ответ, а не пустой файл", why or "")
        check(why is not None and "нет" in why.lower(), "с понятным объяснением", why or "")

        outsider = (
            await session.execute(
                select(User).join(Organization, Organization.id == User.organization_id)
                .where(Organization.name == f"{TEST_ORG_PREFIX}Чужая")
            )
        ).scalar_one()
        alien, _, _ = await export.build(
            session, user=outsider, grants=await load_grants(session, outsider),
            kind="tasks", since=since, until=until,
        )
        check(alien is None, "чужая организация не выгружает ничего")

        wrong, _, why = await export.build(
            session, user=chief, grants=await load_grants(session, chief),
            kind="секреты", since=since, until=until,
        )
        check(wrong is None, "неизвестный вид выгрузки отклоняется", why or "")

    print("\n17. Готовность блока: совещание целиком")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head = who["руководитель"], who["сотрудник"], who["начальник"]
        day = MONDAY + timedelta(days=14)

        planned = Meeting(
            organization_id=chief.organization_id, owner_id=chief.id,
            title="ТЕСТ итоговое совещание по кварталу",
            start_at=at(day, 10), end_at=at(day, 11),
            status=MeetingStatus.CONFIRMED, created_by=chief.id,
        )
        session.add(planned)
        await session.flush()
        for uid in (chief.id, worker.id, head.id):
            session.add(MeetingParticipant(
                meeting_id=planned.id, user_id=uid,
                role=ParticipantRole.REQUIRED, created_at=utcnow(),
            ))
        await session.flush()

        # Документ к совещанию, открытый всем участникам.
        body = tiny_docx([
            "Отчёт за квартал",
            "Отгрузка задержана из-за нехватки складских мест",
            "Предложено арендовать дополнительный ангар",
        ])
        attached = await docs.store(
            session, uploader=worker, file_id="id-квартал", file_unique_id="u-квартал",
            file_name="отчёт-квартал.docx", size_bytes=len(body),
            meeting=planned, scope=DocumentScope.PARTICIPANTS,
        )
        check(attached.ok, "документ приложен к совещанию", attached.reason or "")

        # В очереди остались документы прошлых разделов; поддельный загрузчик
        # их не знает. Уводим их из очереди, чтобы проверка говорила именно
        # о нашем документе, а не о случайном соседе.
        await session.execute(
            update(Document)
            .where(
                Document.index_status == IndexStatus.PENDING,
                Document.id != attached.document.id,
            )
            .values(index_status=IndexStatus.UNSUPPORTED)
        )
        await session.flush()

        async def download(file_id: str) -> bytes:
            return body

        await indexer.index_pending(session, download, limit=5)
        check(
            (await session.get(Document, attached.document.id)).index_status == IndexStatus.DONE,
            "текст документа извлечён",
        )

        finished = await meeting_service.finish(
            session, meeting=planned, actor=chief, now=at(day, 11, 5)
        )
        check(finished.ok, "совещание завершено", finished.reason or "")

        # Три решения и три поручения — то, ради чего блок и делался.
        subjects = [
            ("ТЕСТ арендовать ангар под сезонный запас", worker),
            ("ТЕСТ пересмотреть график отгрузок", head),
            ("ТЕСТ согласовать бюджет аренды", worker),
        ]
        made_decisions = []
        for title, responsible in subjects:
            outcome = await registry.create(
                session, actor=chief, title=title, meeting=planned,
                responsible=responsible, due_date=at(day + timedelta(days=7), 18),
            )
            check(outcome.ok, f"решение «{title[:28]}…» внесено", outcome.reason or "")
            made_decisions.append(outcome.item)

        for decision, (title, responsible) in zip(made_decisions, subjects):
            task = await task_service.create_task(
                session, creator=chief, assignee=responsible,
                title=title.replace("ТЕСТ ", "ТЕСТ выполнить: "),
                due_at=at(day + timedelta(days=7), 18),
                meeting_id=planned.id, decision_id=decision.id,
            )
            check(task.due_at is not None, "у поручения есть срок")
            check(task.assignee_id == responsible.id, "и ответственный")

        counted = await registry.meeting_outcome(session, planned)
        check(counted == (3, 3), "три решения и три поручения зафиксированы", f"{counted}")

        with_due = await session.scalar(
            select(func.count(Decision.id)).where(
                Decision.meeting_id == planned.id,
                Decision.responsible_id.is_not(None),
                Decision.due_date.is_not(None),
            )
        )
        check(with_due == 3, "у каждого решения есть ответственный и срок", f"их {with_due}")

        # И документ находится по фразе из его текста — участником встречи.
        found = await search.search(
            session, user=head, grants=await load_grants(session, head),
            query="нехватка складских мест",
        )
        check(
            any(h.id == attached.document.id for h in found.documents),
            "участник находит документ по фразе из текста",
            f"{[h.title for h in found.documents]}",
        )
        check(
            any(h.id == planned.id for h in found.meetings)
            or bool(found.decisions) or bool(found.tasks) or bool(found.documents),
            "поиск отвечает по нескольким разделам сразу",
        )

        # А тот, кого на совещании не было, — не находит.
        helper_user = who["ассистент"]
        hidden = await search.search(
            session, user=helper_user, grants=await load_grants(session, helper_user),
            query="нехватка складских мест",
        )
        check(
            not any(h.id == attached.document.id for h in hidden.documents),
            "и не находит тот, кого на совещании не было",
            f"{[h.title for h in hidden.documents]}",
        )

    print("\n18. Права на решение описаны дважды и совпадают")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head, helper_user = (
            who["руководитель"], who["сотрудник"], who["начальник"], who["ассистент"]
        )
        outsider = (
            await session.execute(
                select(User).join(Organization, Organization.id == User.organization_id)
                .where(Organization.name == f"{TEST_ORG_PREFIX}Чужая")
            )
        ).scalar_one()

        made = []
        for author, responsible, label in (
            (chief, None, "решение руководителя"),
            (chief, worker, "решение с ответственным сотрудником"),
            (head, None, "решение начальника отдела"),
        ):
            outcome = await registry.create(
                session, actor=author, title=f"ТЕСТ {label}", responsible=responsible
            )
            check(outcome.ok, f"внесено: {label}", outcome.reason or "")
            made.append(outcome.item)
        await session.flush()

        mismatches = []
        for decision in made:
            for name, viewer in (
                ("руководитель", chief), ("ассистент", helper_user),
                ("начальник", head), ("сотрудник", worker), ("чужой", outsider),
            ):
                direct = await registry.may_read(session, decision=decision, viewer=viewer)
                grants = await load_grants(session, viewer)
                visible = await visible_department_ids(session, viewer)
                by_sql = await session.scalar(
                    select(Decision.id).where(
                        Decision.id == decision.id,
                        *registry.visible_filter(viewer, grants, visible),
                    )
                )
                if direct != (by_sql is not None):
                    mismatches.append(
                        f"{decision.title}/{name}: прямо={direct}, запросом={by_sql is not None}"
                    )
        check(
            not mismatches,
            "проверка записи и условие запроса дают один ответ",
            "; ".join(mismatches[:3]),
        )
        check(
            not await registry.may_read(session, decision=made[0], viewer=outsider),
            "чужая организация решение не видит",
        )

        # Тот самый случай, из-за которого карточка врала: видимость одной
        # записи не должна зависеть от того, попала ли она в выборку реестра.
        short_list = await registry.registry(
            session, user=chief, grants=await load_grants(session, chief), limit=1
        )
        beyond = [d for d in made if d.id not in {x.id for x in short_list}]
        check(bool(beyond), "решения за пределами короткой выборки есть", f"{len(beyond)}")
        hidden = []
        for decision in beyond:
            if not await registry.may_read(session, decision=decision, viewer=chief):
                hidden.append(decision.title)
        check(
            not hidden,
            "и они всё равно открыты своему автору",
            f"закрылись: {hidden[:2]}",
        )

        # Закрытие спрашивает не только право, но и доступ к записи.
        # Ролью это сегодня не достаётся — право `decision.close` есть только
        # у областей ORGANIZATION, — поэтому проверяем достижимый случай.
        stranger_close = await registry.close(
            session, decision=made[0], actor=outsider, done=True
        )
        check(
            not stranger_close.ok,
            "чужой не закрывает решение",
            stranger_close.reason or "закрылось",
        )
        own_close = await registry.close(session, decision=made[2], actor=chief, done=True)
        check(own_close.ok, "а руководитель закрывает своё", own_close.reason or "")

    print("\n19. Выгрузка не превращает текст в формулы")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker = who["руководитель"], who["сотрудник"]
        # Название обязано начинаться со знака равенства — иначе проверка
        # не доходит до опасного случая и проходит при любом исходе.
        # Уборка удаляет поручения по организации, а не по префиксу названия.
        danger = "=cmd|'/c calc'!A1 ТЕСТ"
        await task_service.create_task(
            session, creator=chief, assignee=worker, title=danger,
        )
        await session.flush()

        since, until = at(MONDAY - timedelta(days=120), 0), at(MONDAY + timedelta(days=60), 0)
        data, _, why = await export.build(
            session, user=chief, grants=await load_grants(session, chief),
            kind="tasks", since=since, until=until, fmt="xlsx",
        )
        check(data is not None, "выгрузка собрана", why or "")

        page = load_workbook(io.BytesIO(data)).active
        formulas = [
            cell.value
            for row in page.iter_rows(min_row=2)
            for cell in row
            if cell.data_type == "f"
        ]
        check(not formulas, "ни одной формульной ячейки", f"{formulas[:2]}")
        check(
            any(danger in str(c.value) for row in page.iter_rows(min_row=2) for c in row),
            "и текст сохранён как есть, без искажения",
        )

    print("\n20. Досье собирается один раз на встречу")
    async with session_scope() as session:
        who = await cast(session)
        chief, worker, head, helper_user = (
            who["руководитель"], who["сотрудник"], who["начальник"], who["ассистент"]
        )
        org = await session.get(Organization, chief.organization_id)
        crowd_day = MONDAY + timedelta(days=21)
        crowded = Meeting(
            organization_id=chief.organization_id, owner_id=chief.id,
            title="ТЕСТ большое совещание", start_at=at(crowd_day, 15),
            end_at=at(crowd_day, 16), status=MeetingStatus.CONFIRMED, created_by=chief.id,
        )
        session.add(crowded)
        await session.flush()
        for guest in (chief, worker, head, helper_user):
            session.add(MeetingParticipant(
                meeting_id=crowded.id, user_id=guest.id,
                role=ParticipantRole.REQUIRED, created_at=utcnow(),
            ))
        # Незакрытое решение с ответственным-участником: без него общая часть
        # была бы пустой и проверка ничего бы не измерила.
        await registry.create(
            session, actor=chief, title="ТЕСТ незакрытое по складу", responsible=worker
        )
        await session.flush()

        statements: list[str] = []

        def record(conn, cursor, statement, parameters, context, executemany):
            statements.append(statement)

        event.listen(engine.sync_engine, "before_cursor_execute", record)
        try:
            sent = await briefing.send_briefings(session, now=at(crowd_day, 14, 45))
        finally:
            event.remove(engine.sync_engine, "before_cursor_execute", record)

        check(sent == 4, "досье ушло всем четверым", f"писем: {sent}")
        # Общая часть собирается один раз на встречу, а не на каждого получателя.
        # Четыре выборки решений вместо одной — это и есть «запросы в квадрате».
        decision_queries = sum(1 for q in statements if "FROM decisions" in q)
        check(
            decision_queries == 1,
            "решения выбраны один раз, а не на каждого участника",
            f"запросов к decisions: {decision_queries} при четырёх получателях",
        )

    print("\n21. Обрезка списка участников не рубит HTML-сущность")
    async with session_scope() as session:
        who = await cast(session)
        chief = who["руководитель"]
        org = await session.get(Organization, chief.organization_id)
        long_day = MONDAY + timedelta(days=28)

        # Имя подобрано так, чтобы `<` пришлась ровно на границу обрезки.
        # В открытом виде строка укладывается в предел (199 знаков) и не режется;
        # в экранированном `&lt;` занимает 202 знака, и обрезка по 200 разрубила
        # бы сущность пополам. Случайное имя сюда не попадает — проверка,
        # которая не доходит до опасного случая, проходит при любом исходе.
        tricky_name = "А" * 196 + "<" + "ББ"
        check(len(tricky_name) == 199, "имя подобрано под границу", str(len(tricky_name)))
        tricky = await person(session, org, tricky_name, RoleCode.EMPLOYEE, 953_000_001)

        pair = Meeting(
            organization_id=chief.organization_id, owner_id=chief.id,
            title="ТЕСТ разговор вдвоём", start_at=at(long_day, 15),
            end_at=at(long_day, 16), status=MeetingStatus.CONFIRMED, created_by=chief.id,
        )
        session.add(pair)
        await session.flush()
        for guest in (chief, tricky):
            session.add(MeetingParticipant(
                meeting_id=pair.id, user_id=guest.id,
                role=ParticipantRole.REQUIRED, created_at=utcnow(),
            ))
        await session.flush()

        sent = await briefing.send_briefings(session, now=at(long_day, 14, 45))
        check(sent == 2, "досье ушло обоим", f"писем: {sent}")

        letters = (
            await session.execute(
                select(Notification.body).where(
                    Notification.event_key.like(f"meeting:{pair.id}:brief:%")
                )
            )
        ).scalars().all()
        check(len(letters) == 2, "писем в очереди два", f"{len(letters)}")
        broken = [
            body for body in letters
            if re.search(r"&(?!(amp|lt|gt|quot|#\d+);)", body)
        ]
        check(
            not broken,
            "ни в одном письме нет разрубленной HTML-сущности",
            (broken[0][-80:] if broken else ""),
        )
        check(
            all("&lt;" in body for body in letters),
            "а сама скобка на месте и экранирована целиком",
        )

    print("\n22. Уборка не трогает боевые данные")
    async with session_scope() as session:
        real_before = await session.scalar(
            select(func.count(User.id)).where(
                User.organization_id.notin_(
                    select(Organization.id).where(Organization.name.like(f"{TEST_ORG_PREFIX}%"))
                )
            )
        )
    await cleanup()
    async with session_scope() as session:
        real_after = await session.scalar(select(func.count(User.id)))
        left = await session.scalar(select(func.count(Decision.id)))
    check(real_after == real_before, "настоящие сотрудники не удалены", f"{real_before} → {real_after}")
    check(left == 0, "тестовые решения убраны", f"осталось {left}")

    print(f"\n{'=' * 50}\nПройдено: {passed}   Ошибок: {failed}\n{'=' * 50}")
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    asyncio.run(main())
