"""Извлечение текста из документов и запись его в поисковый индекс.

Работа фоновая по существу: разбор PDF на двести страниц занимает секунды,
а бот обязан ответить человеку сразу. Поэтому загрузка только принимает файл
и ставит отметку «ждёт разбора», а сюда приходит отдельный цикл.

Сбой разбора не теряет документ. Битый файл, скан без текстового слоя,
незнакомый формат — всё это состояния документа, а не причины его потерять:
он остаётся доступен и находится по имени.
"""
import io
import logging
import os

from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.timeutil import utcnow
from app.models import Document, DocumentText, IndexStatus

log = logging.getLogger("seta.indexer")

# Столько текста храним и индексируем. Ограничение не наше: tsvector в Postgres
# не может быть больше мегабайта, а документ на тысячу страниц столько и даст.
# Первые полмиллиона знаков — это примерно двести страниц; если искомая фраза
# не встретилась там ни разу, поиск по документу всё равно не тот инструмент.
MAX_TEXT_CHARS = 500_000

# Управляющие символы, которые база не примет. Нулевой байт в тексте — обычное
# дело: достаточно, чтобы двоичный файл прислали под именем «отчёт.txt».
# Табуляцию, перевод строки и возврат каретки оставляем.
_CONTROL = {code: None for code in range(32) if code not in (9, 10, 13)}
_CONTROL[127] = None


def clean(text: str) -> str:
    """Убирает управляющие символы: PostgreSQL отвергает строку с 0x00 целиком."""
    return (text or "").translate(_CONTROL)


def extract(data: bytes, file_name: str) -> tuple[str, int]:
    """Достаёт текст из файла. Возвращает (текст, число страниц или листов).

    Бросает исключение при непригодном файле — вызывающий превращает это
    в состояние документа, а не в потерю.
    """
    extension = os.path.splitext((file_name or "").lower())[1]

    if extension in (".txt", ".md", ".csv"):
        for encoding in ("utf-8", "cp1251"):
            try:
                return data.decode(encoding), 1
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace"), 1

    if extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
            if sum(len(p) for p in parts) > MAX_TEXT_CHARS:
                break
        return "\n".join(parts), len(reader.pages)

    if extension == ".docx":
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # Таблицы в договорах несут половину смысла: без них поиск по документу
        # находил бы преамбулу и терял сумму, сроки и предмет.
        for table in document.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts), 1

    if extension == ".xlsx":
        from openpyxl import load_workbook

        book = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in book.worksheets:
            parts.append(sheet.title)
            for row in sheet.iter_rows(values_only=True):
                parts.append(" ".join(str(v) for v in row if v is not None))
                if sum(len(p) for p in parts) > MAX_TEXT_CHARS:
                    break
        sheets = len(book.worksheets)
        book.close()
        return "\n".join(parts), sheets

    raise ValueError(f"формат {extension or 'без расширения'} не разбирается")


async def save_text(
    session: AsyncSession, *, document: Document, content: str, pages: int
) -> None:
    """Записывает извлечённый текст и сразу считает поисковый вектор.

    Вектор хранится, а не считается на каждый запрос: морфология по документу
    на двести страниц — заметная работа, и делать её при каждом поиске незачем.
    """
    content = clean(content).strip()[:MAX_TEXT_CHARS]
    existing = (
        await session.execute(
            select(DocumentText).where(DocumentText.document_id == document.id)
        )
    ).scalar_one_or_none()

    if existing is None:
        existing = DocumentText(document_id=document.id, extracted_at=utcnow())
        session.add(existing)
    existing.content = content
    existing.pages = max(0, pages)
    existing.extracted_at = utcnow()
    existing.search_vector = func.to_tsvector("russian", content)
    await session.flush()


async def index_pending(session: AsyncSession, download, limit: int = 5) -> dict[str, int]:
    """Разбирает документы, ждущие очереди.

    `download(file_id) -> bytes` передаётся снаружи: службе незачем знать,
    что файлы лежат в Telegram, а проверкам — обращаться в сеть.
    """
    waiting = (
        await session.execute(
            select(Document)
            .where(Document.index_status == IndexStatus.PENDING)
            .order_by(Document.created_at)
            .limit(limit)
        )
    ).scalars().all()

    stats = {"done": 0, "empty": 0, "failed": 0}
    for document in waiting:
        try:
            data = await download(document.file_id)
            content, pages = extract(data, document.file_name)
        except Exception as error:  # noqa: BLE001 — любой сбой разбора это состояние
            document.index_status = IndexStatus.FAILED
            document.index_error = f"{type(error).__name__}: {error}"[:300]
            stats["failed"] += 1
            log.warning("не разобран документ %s: %s", document.id, error)
            continue

        if not clean(content).strip():
            # Скан без текстового слоя. Не ошибка: документ цел, просто
            # искать по нему можно только по имени.
            document.index_status = IndexStatus.EMPTY
            stats["empty"] += 1
            continue

        # Запись под точкой сохранения: отказ базы на одном документе не должен
        # обрывать проход и уж тем более откатывать разобранные до него.
        # Без этого проход падал бы на первом же файле с нулевым байтом внутри
        # и повторял бы падение каждую минуту, не двигаясь дальше.
        savepoint = await session.begin_nested()
        try:
            await save_text(session, document=document, content=content, pages=pages)
        except Exception as error:  # noqa: BLE001 — отказ записи это тоже состояние
            await savepoint.rollback()
            document.index_status = IndexStatus.FAILED
            document.index_error = f"{type(error).__name__}: {error}"[:300]
            stats["failed"] += 1
            log.warning("не записан текст документа %s: %s", document.id, error)
            continue

        document.index_status = IndexStatus.DONE
        document.index_error = None
        stats["done"] += 1

    await session.flush()
    return stats
